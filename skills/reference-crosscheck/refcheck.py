#!/usr/bin/env python3
"""
refcheck.py - Universal reference cross-check engine for HWP / HWPX / Word.

Tier 1 (always, offline, deterministic):
  - Extract location-tagged text from .docx / .hwpx / .hwp.
  - Auto-detect citation style (numeric bracket, numeric superscript, author-year).
  - Parse in-text citations and the reference list.
  - Detect:
      * phantom   : cited in text but missing from the reference list
      * orphan    : in the reference list but never cited
      * numbering : cited number exceeds list size / list numbering has gaps
      * duplicate : the same source listed more than once
      * multiloc  : the same reference present in >1 location
                    (body list + footnote/endnote, or >1 reference list)
  - Emit a structured markdown + JSON findings report.
  - Produce an annotated copy of the source (graceful per-format degradation).

Tier 2 (opt-in, --tier2-queue): emits crossref_queue.json of reference
entries to be verified for real-world existence. The verification itself is
done by the paper-ref-hunter subagent (Crossref/DOI) orchestrated by SKILL.md,
not by this script (keeps Tier 1 offline-safe).

Usage:
  python refcheck.py INPUT [--outdir DIR] [--style auto|numeric|author_year]
                     [--no-annotate] [--tier2-queue]
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass, field, asdict
from difflib import SequenceMatcher

# --------------------------------------------------------------------------- #
#  Data model
# --------------------------------------------------------------------------- #


@dataclass
class Segment:
    """One contiguous text unit, tagged by where it lives in the document."""

    loc: str  # body | footnote | endnote | table | header | preview
    text: str
    ref_idx: int = -1  # paragraph index for docx body (annotation anchor)


@dataclass
class Citation:
    raw: str  # as it appears, e.g. "[12, 14-16]"
    key: str  # normalized single key: "12" or "kim2020"
    loc: str
    seg: int


@dataclass
class RefEntry:
    num: int  # 1-based list position (numeric style) or -1
    key: str  # "12" or "kim2020"
    raw: str
    sig: str  # dedup signature
    doi: str = ""
    first_author: str = ""
    year: str = ""
    title: str = ""
    locs: list = field(default_factory=list)  # locations this entry was seen in


@dataclass
class Findings:
    fmt: str = ""
    style: str = ""
    n_citations: int = 0
    n_entries: int = 0
    ref_section_locations: list = field(default_factory=list)
    phantom: list = field(default_factory=list)
    orphan: list = field(default_factory=list)
    numbering: list = field(default_factory=list)
    duplicate: list = field(default_factory=list)
    multiloc: list = field(default_factory=list)
    notes: list = field(default_factory=list)


# --------------------------------------------------------------------------- #
#  Extraction layer (format-dispatched)
# --------------------------------------------------------------------------- #

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_docx(path):
    """Body (with superscript flag), tables, footnotes, endnotes."""
    from docx import Document

    doc = Document(path)
    segs = []

    for i, p in enumerate(doc.paragraphs):
        txt = p.text
        # Reconstruct superscript citation tokens from runs.
        sup = ""
        for r in p.runs:
            if r.font.superscript and re.fullmatch(r"[\d,\s\-–]+", r.text or ""):
                sup += r.text
        if sup.strip():
            txt = f"{txt} ⸺{sup.strip()}⸺"  # mark superscripts
        if txt.strip():
            segs.append(Segment("body", txt, ref_idx=i))

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    segs.append(Segment("table", ct))

    # Footnotes / endnotes are not exposed by python-docx -> read the parts.
    with zipfile.ZipFile(path) as z:
        for part, loc in (("word/footnotes.xml", "footnote"),
                          ("word/endnotes.xml", "endnote")):
            if part not in z.namelist():
                continue
            import xml.etree.ElementTree as ET

            root = ET.fromstring(z.read(part))
            tag = "footnote" if loc == "footnote" else "endnote"
            for note in root.findall(f"{W_NS}{tag}"):
                nid = note.get(f"{W_NS}id")
                if nid in ("-1", "0"):  # separator / continuation
                    continue
                txt = "".join(t.text or "" for t in note.iter(f"{W_NS}t"))
                if txt.strip():
                    segs.append(Segment(loc, txt))
    return "docx", segs


def _hwpx_text_from_section(xml_bytes):
    """Collect <hp:t> text and tag footnote/endnote regions for HWPX/OWPML.

    Walk recursively while tracking the enclosing location so that text
    inside a footNote/endNote is NOT also counted as body text.
    """
    import xml.etree.ElementTree as ET

    out = []  # (loc, text)

    def walk(el, loc):
        tag = el.tag.split("}")[-1]
        if tag in ("footNote", "footnote"):
            loc = "footnote"
        elif tag in ("endNote", "endnote"):
            loc = "endnote"
        if tag == "t" and el.text:
            out.append((loc, el.text))
        for child in el:
            walk(child, loc)

    walk(ET.fromstring(xml_bytes), "body")

    # Merge consecutive footnote/endnote fragments into one segment each.
    merged = []
    for loc, txt in out:
        if merged and merged[-1][0] == loc and loc in ("footnote", "endnote"):
            merged[-1] = (loc, merged[-1][1] + txt)
        else:
            merged.append((loc, txt))
    return merged


def extract_hwpx(path):
    segs = []
    with zipfile.ZipFile(path) as z:
        sections = sorted(n for n in z.namelist()
                          if re.match(r"Contents/section\d+\.xml", n))
        for s in sections:
            for loc, txt in _hwpx_text_from_section(z.read(s)):
                if txt.strip():
                    segs.append(Segment(loc, txt))
    return "hwpx", segs


def extract_hwp(path):
    """Binary HWP 5.0: hwp5txt CLI -> TextTransform -> olefile preview."""
    # 1) hwp5txt CLI (most reliable)
    exe = shutil.which("hwp5txt")
    if exe:
        try:
            out = subprocess.run([exe, path], capture_output=True,
                                  timeout=120)
            text = out.stdout.decode("utf-8", "replace")
            if text.strip():
                segs = [Segment("body", ln) for ln in text.splitlines()
                        if ln.strip()]
                return "hwp", segs
        except Exception:
            pass
    # 2) hwp5 python TextTransform
    try:
        import io
        from hwp5.hwp5txt import TextTransform
        from hwp5.xmlmodel import Hwp5File

        buf = io.StringIO()
        TextTransform().transform_hwp5_to_text(Hwp5File(path), buf)
        text = buf.getvalue()
        if text.strip():
            segs = [Segment("body", ln) for ln in text.splitlines()
                    if ln.strip()]
            return "hwp", segs
    except Exception:
        pass
    # 3) olefile PrvText (preview only - truncated; warn the caller)
    import olefile

    ole = olefile.OleFileIO(path)
    if ole.exists("PrvText"):
        raw = ole.openstream("PrvText").read()
        text = raw.decode("utf-16-le", "replace")
        segs = [Segment("preview", ln) for ln in text.splitlines()
                if ln.strip()]
        return "hwp-preview", segs
    raise RuntimeError("HWP text extraction failed (no hwp5 / no PrvText).")


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".hwpx":
        return extract_hwpx(path)
    if ext == ".hwp":
        return extract_hwp(path)
    raise ValueError(f"Unsupported format: {ext} (use .docx/.hwpx/.hwp)")


# --------------------------------------------------------------------------- #
#  Citation + reference parsing (format-independent)
# --------------------------------------------------------------------------- #

REF_HEADING = re.compile(
    r"^\s*(?:\d+\.?\s*)?"
    r"(references?|bibliography|works\s+cited|literature\s+cited|"
    r"참고\s*문헌|인용\s*문헌|참고자료)\s*$",
    re.IGNORECASE,
)
BRACKET_CITE = re.compile(r"\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]")
SUP_CITE = re.compile(r"⸺([\d,\s\-–]+)⸺")
AY_CITE = re.compile(r"\(([^()]*?(?:1[89]|20)\d{2}[a-z]?[^()]*?)\)")
DOI_RE = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+")


def expand_numeric(token):
    """'[12, 14-16]' -> ['12','14','15','16']."""
    nums = []
    body = token.strip("[] ")
    for part in re.split(r"\s*,\s*", body):
        m = re.match(r"(\d+)\s*[-–]\s*(\d+)$", part)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            nums += [str(n) for n in range(min(a, b), max(a, b) + 1)]
        elif part.strip().isdigit():
            nums.append(part.strip())
    return nums


def ay_key(text):
    """First-author surname + year -> 'kim2020' / '홍2020'."""
    ym = re.search(r"(1[89]|20)\d{2}", text)
    year = ym.group(0) if ym else ""
    head = re.split(r"[,;]| and | & | 외| et al", text.strip())[0].strip()
    surname = re.split(r"\s+", head)[-1] if head else head
    surname = re.sub(r"[^\w가-힣]", "", surname).lower()
    return f"{surname}{year}" if surname or year else ""


def detect_style(segs):
    body = " ".join(s.text for s in segs if s.loc in ("body", "preview"))
    nb = len(BRACKET_CITE.findall(body)) + len(SUP_CITE.findall(body))
    na = len(AY_CITE.findall(body))
    return "numeric" if nb >= na else "author_year"


def split_ref_sections(segs):
    """Return list of (heading_seg_index). Multiple => list in >1 location."""
    idxs = [i for i, s in enumerate(segs) if REF_HEADING.match(s.text.strip())]
    return idxs


def parse_entries(segs, start, style):
    """Parse reference entries from `start`+1 until the segment list ends."""
    entries = []
    buf_num, buf_txt = None, []

    def flush(num, txt):
        txt = txt.strip()
        if not txt:
            return
        doi_m = DOI_RE.search(txt)
        doi = doi_m.group(0).rstrip(".") if doi_m else ""
        if style == "numeric":
            key = str(num)
        else:
            key = ay_key(txt)
        ym = re.search(r"(1[89]|20)\d{2}", txt)
        sig = doi.lower() if doi else re.sub(
            r"[^\w가-힣]", "", txt.lower())[:80]
        entries.append(RefEntry(
            num=num if num is not None else -1, key=key, raw=txt, sig=sig,
            doi=doi, first_author=re.split(r"[,\.]", txt)[0].strip()[:40],
            year=ym.group(0) if ym else "",
            title="", locs=[segs[start].loc]))

    seq = 0
    for s in segs[start + 1:]:
        line = s.text.strip()
        if not line:
            continue
        if REF_HEADING.match(line):  # next ref section -> stop this block
            break
        m = re.match(r"^\[?\(?(\d{1,4})[\]\).·]\s+(.*)$", line)
        if style == "numeric" and m:
            if buf_txt:
                flush(buf_num, " ".join(buf_txt))
            buf_num = int(m.group(1))
            buf_txt = [m.group(2)]
        elif style == "numeric" and buf_num is not None:
            buf_txt.append(line)  # wrapped continuation
        else:  # author-year: one entry per paragraph
            seq += 1
            flush(seq, line)
    if buf_txt:
        flush(buf_num, " ".join(buf_txt))
    return entries


def parse_citations(segs, ref_starts, style):
    cites = []
    cutoff = min(ref_starts) if ref_starts else len(segs)
    for i, s in enumerate(segs):
        if i >= cutoff and s.loc in ("body", "table"):
            continue  # don't read the reference list itself as in-text cites
        if style == "numeric":
            for m in list(BRACKET_CITE.finditer(s.text)):
                for n in expand_numeric(m.group(0)):
                    cites.append(Citation(m.group(0), n, s.loc, i))
            for m in SUP_CITE.finditer(s.text):
                for n in expand_numeric(m.group(1)):
                    cites.append(Citation(m.group(0), n, s.loc, i))
        else:
            for m in AY_CITE.finditer(s.text):
                for part in re.split(r"\s*;\s*", m.group(1)):
                    k = ay_key(part)
                    if k:
                        cites.append(Citation(m.group(0), k, s.loc, i))
    return cites


# --------------------------------------------------------------------------- #
#  Cross-check engine
# --------------------------------------------------------------------------- #


def crosscheck(segs, fmt, style):
    f = Findings(fmt=fmt, style=style)
    ref_starts = split_ref_sections(segs)
    f.ref_section_locations = [
        {"seg": i, "loc": segs[i].loc, "heading": segs[i].text.strip()}
        for i in ref_starts
    ]
    if len(ref_starts) > 1:
        f.notes.append(
            f"Reference-list heading found in {len(ref_starts)} places "
            f"-> potential duplicated/per-section list. Using the largest.")

    # Pick the section that yields the most entries as the primary list;
    # parse every section so we can detect the same ref in >1 location.
    all_blocks = []
    for st in ref_starts:
        all_blocks.append((st, parse_entries(segs, st, style)))
    if all_blocks:
        primary_start, entries = max(all_blocks, key=lambda b: len(b[1]))
    else:
        primary_start, entries = len(segs), []
        f.notes.append("No reference-list heading detected "
                        "(References/참고문헌 ...). Phantom check skipped.")

    cites = parse_citations(segs, ref_starts, style)
    f.n_citations = len(cites)
    f.n_entries = len(entries)

    entry_keys = {e.key for e in entries}
    cited_keys = {c.key for c in cites}

    # ---- phantom: cited but no entry ----
    for c in sorted({c.key for c in cites}):
        if c not in entry_keys:
            if style == "author_year":
                # fuzzy: tolerate minor surname spelling drift
                if any(SequenceMatcher(None, c, e).ratio() > 0.9
                       for e in entry_keys):
                    continue
            occ = [s for s in cites if s.key == c]
            f.phantom.append({
                "key": c, "count": len(occ),
                "raw": occ[0].raw,
                "where": sorted({o.loc for o in occ})})

    # ---- orphan: entry never cited ----
    for e in entries:
        if e.key not in cited_keys:
            f.orphan.append({"num": e.num, "key": e.key,
                             "ref": e.raw[:160]})

    # ---- numbering sanity (numeric only) ----
    if style == "numeric" and entries:
        max_num = max((e.num for e in entries if e.num > 0), default=0)
        cited_nums = sorted({int(c.key) for c in cites if c.key.isdigit()})
        over = [n for n in cited_nums if n > max_num]
        if over:
            f.numbering.append(
                f"Cited numbers exceed reference-list size "
                f"(list has {max_num} entries, cited up to {max(over)}): "
                f"{over}")
        nums = sorted(e.num for e in entries if e.num > 0)
        gaps = [n for n in range(1, max_num + 1) if n not in nums]
        if gaps:
            f.numbering.append(f"Reference list numbering gaps: {gaps}")
        dups_n = sorted({n for n in nums if nums.count(n) > 1})
        if dups_n:
            f.numbering.append(
                f"Reference list reuses the same number: {dups_n}")

    # ---- duplicate entries (same source listed twice) ----
    seen = {}
    for e in entries:
        seen.setdefault(e.sig, []).append(e)
    for sig, grp in seen.items():
        if len(grp) > 1 and sig:
            f.duplicate.append({
                "signature": sig[:60],
                "members": [{"num": g.num, "ref": g.raw[:140]} for g in grp]})
    # near-duplicate (different formatting, same paper)
    sigs = [(e.sig, e) for e in entries if e.sig]
    for a in range(len(sigs)):
        for b in range(a + 1, len(sigs)):
            if sigs[a][0] == sigs[b][0]:
                continue
            r = SequenceMatcher(None, sigs[a][0], sigs[b][0]).ratio()
            if r > 0.92:
                f.duplicate.append({
                    "signature": f"~near {r:.2f}",
                    "members": [
                        {"num": sigs[a][1].num, "ref": sigs[a][1].raw[:140]},
                        {"num": sigs[b][1].num, "ref": sigs[b][1].raw[:140]}]})

    # ---- same reference present in multiple locations ----
    # (e.g. listed in the end bibliography AND spelled out in a footnote,
    #  or appearing in more than one reference-list block)
    loc_index = {}
    for st, blk in all_blocks:
        for e in blk:
            loc_index.setdefault(e.sig, set()).add(segs[st].loc)
    # footnotes/endnotes that themselves look like full references
    for s in segs:
        if s.loc in ("footnote", "endnote") and DOI_RE.search(s.text):
            sig = (DOI_RE.search(s.text).group(0).lower())
            loc_index.setdefault(sig, set()).add(s.loc)
    for sig, locs in loc_index.items():
        if len(locs) > 1 and sig:
            ref_txt = next((e.raw[:140] for _, blk in all_blocks
                            for e in blk if e.sig == sig), sig[:60])
            f.multiloc.append({"signature": sig[:60],
                               "locations": sorted(locs),
                               "ref": ref_txt})

    return f, entries, cites


# --------------------------------------------------------------------------- #
#  Tier 2 queue (real-existence verification handled by paper-ref-hunter)
# --------------------------------------------------------------------------- #


def build_tier2_queue(entries, findings):
    """Entries worth a Crossref check: no DOI, orphan, or duplicate."""
    dup_nums = {m["num"] for d in findings.duplicate for m in d["members"]}
    orphan_nums = {o["num"] for o in findings.orphan}
    queue = []
    for e in entries:
        reasons = []
        if not e.doi:
            reasons.append("no_doi_in_text")
        if e.num in dup_nums:
            reasons.append("duplicate_candidate")
        if e.num in orphan_nums:
            reasons.append("orphan")
        queue.append({
            "num": e.num, "key": e.key, "doi": e.doi,
            "first_author": e.first_author, "year": e.year,
            "raw": e.raw, "flag_reasons": reasons,
            "priority": "high" if reasons else "low"})
    return queue


# --------------------------------------------------------------------------- #
#  Annotated-copy output (graceful per-format degradation)
# --------------------------------------------------------------------------- #


def _annotation_targets(findings):
    """Plain-substring needles to flag in the copy, with their note.

    Covers every finding category, not just phantom:
      - phantom   : the in-text citation string ('[99]' / '(Park, 2099)')
      - duplicate : a stable prefix of each duplicated reference entry
      - orphan    : a stable prefix of each uncited reference entry
    Plain substrings (not regex) so `in` / `str.replace` stay safe.
    """
    targets = []  # (needle, note)
    for p in findings.phantom:
        targets.append((p["raw"],
                         f"PHANTOM: cited but not in reference list"))
    for d in findings.duplicate:
        nums = ", ".join(f"#{m['num']}" for m in d["members"])
        for m in d["members"]:
            needle = re.sub(r"^\s*\[?\(?\d+[\]\).·]?\s*", "",
                            m["ref"]).strip()[:34]
            if needle:
                targets.append(
                    (needle, f"DUPLICATE reference entry ({nums}, "
                             f"sig {d['signature']})"))
    for o in findings.orphan:
        needle = re.sub(r"^\s*\[?\(?\d+[\]\).·]?\s*", "",
                        o["ref"]).strip()[:34]
        if needle:
            targets.append(
                (needle, f"ORPHAN: reference #{o['num']} never cited"))
    # Deduplicate identical needles, keeping the first note.
    seen, uniq = set(), []
    for needle, note in targets:
        if needle and needle not in seen:
            seen.add(needle)
            uniq.append((needle, note))
    return uniq


def annotate_docx(src, out, findings):
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    shutil.copyfile(src, out)
    doc = Document(out)
    targets = _annotation_targets(findings)
    touched = 0

    paras = list(doc.paragraphs)
    for t in doc.tables:  # reference lists are sometimes in tables
        for row in t.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)

    for p in paras:
        ptxt = p.text
        hits = [(needle, note) for needle, note in targets if needle in ptxt]
        if not hits or not p.runs:
            continue
        msg = "; ".join(sorted({note for _, note in hits}))
        try:
            doc.add_comment(p.runs, text=f"[refcheck] {msg}",
                            author="reference-crosscheck", initials="RC")
        except Exception:
            mark = p.add_run(f"  〔refcheck: {msg}〕")
            mark.font.highlight_color = WD_COLOR_INDEX.YELLOW
        for r in p.runs:
            if any(needle in (r.text or "") for needle, _ in hits):
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        touched += 1

    doc.save(out)
    return touched


def annotate_hwpx(src, out, findings):
    """Best-effort: inject a visible marker after each flagged string."""
    targets = _annotation_targets(findings)
    with zipfile.ZipFile(src) as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    touched = 0
    for n in list(data):
        if not re.match(r"Contents/section\d+\.xml", n):
            continue
        xml = data[n].decode("utf-8", "replace")
        for needle, note in targets:
            if needle in xml:
                xml = xml.replace(
                    needle, f"{needle}〔refcheck: {note}〕", 1)
                touched += 1
        data[n] = xml.encode("utf-8")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    return touched


def annotate_hwp_fallback(segs, out, findings):
    """Binary .hwp cannot be safely rewritten -> annotated text sidecar."""
    targets = _annotation_targets(findings)
    lines = ["# refcheck annotated text (source .hwp not modifiable)",
             "# 〔refcheck: ...〕 appended to any line containing a flagged "
             "string", ""]
    touched = 0
    for s in segs:
        notes = [note for needle, note in targets if needle in s.text]
        if notes:
            touched += len(notes)
            lines.append(f"[{s.loc}] {s.text}   〔refcheck: "
                          + " | ".join(sorted(set(notes))) + "〕")
        else:
            lines.append(f"[{s.loc}] {s.text}")
    with open(out, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    return touched


# --------------------------------------------------------------------------- #
#  Reporting
# --------------------------------------------------------------------------- #


def write_markdown(f: Findings, path):
    L = []
    L.append(f"# Reference cross-check report\n")
    L.append(f"- Source format: `{f.fmt}`")
    L.append(f"- Citation style: `{f.style}`")
    L.append(f"- In-text citations parsed: {f.n_citations}")
    L.append(f"- Reference-list entries: {f.n_entries}")
    if f.ref_section_locations:
        locs = ", ".join(
            f"{r['heading']!r}@seg{r['seg']}({r['loc']})"
            for r in f.ref_section_locations)
        L.append(f"- Reference-list heading(s): {locs}")
    L.append("")

    def block(title, items, render):
        L.append(f"## {title} ({len(items)})")
        if not items:
            L.append("- none\n")
            return
        for it in items:
            L.append(render(it))
        L.append("")

    block("Phantom (cited, not in reference list)", f.phantom,
          lambda p: f"- **{p['key']}** x{p['count']} "
                    f"raw=`{p['raw']}` in {p['where']}")
    block("Orphan (in list, never cited)", f.orphan,
          lambda o: f"- #{o['num']} `{o['key']}` - {o['ref']}")
    block("Duplicate reference entries", f.duplicate,
          lambda d: f"- {d['signature']}: "
                    + " | ".join(f"#{m['num']} {m['ref']}"
                                 for m in d["members"]))
    block("Same reference in multiple locations", f.multiloc,
          lambda m: f"- {m['locations']} :: {m['ref']}")
    L.append("## Numbering sanity")
    L += ([f"- {n}" for n in f.numbering] or ["- ok"])
    L.append("")
    if f.notes:
        L.append("## Notes")
        L += [f"- {n}" for n in f.notes]
        L.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(L))


# --------------------------------------------------------------------------- #
#  CLI
# --------------------------------------------------------------------------- #


def main(argv=None):
    ap = argparse.ArgumentParser(description="Universal reference cross-check")
    ap.add_argument("input")
    ap.add_argument("--outdir", default=None)
    ap.add_argument("--style", choices=["auto", "numeric", "author_year"],
                    default="auto")
    ap.add_argument("--no-annotate", action="store_true")
    ap.add_argument("--tier2-queue", action="store_true",
                    help="emit crossref_queue.json for paper-ref-hunter")
    args = ap.parse_args(argv)

    src = os.path.abspath(args.input)
    if not os.path.isfile(src):
        print(f"ERROR: not found: {src}", file=sys.stderr)
        return 2
    outdir = os.path.abspath(args.outdir) if args.outdir \
        else os.path.dirname(src)
    os.makedirs(outdir, exist_ok=True)
    base = os.path.basename(src)
    name_noext, ext = os.path.splitext(base)
    # Report stem carries the source extension so checking same-named
    # .docx and .hwp into one folder never collides.
    stem = f"{name_noext}{ext.replace('.', '_')}"

    fmt, segs = extract(src)
    style = args.style if args.style != "auto" else detect_style(segs)
    findings, entries, _ = crosscheck(segs, fmt, style)

    md_path = os.path.join(outdir, f"{stem}.refcheck.md")
    json_path = os.path.join(outdir, f"{stem}.refcheck.json")
    write_markdown(findings, md_path)
    with open(json_path, "w", encoding="utf-8") as fh:
        json.dump(asdict(findings), fh, ensure_ascii=False, indent=2)

    produced = [md_path, json_path]

    if args.tier2_queue:
        q = build_tier2_queue(entries, findings)
        qp = os.path.join(outdir, f"{stem}.crossref_queue.json")
        with open(qp, "w", encoding="utf-8") as fh:
            json.dump(q, fh, ensure_ascii=False, indent=2)
        produced.append(qp)

    if not args.no_annotate:
        if fmt == "docx":
            ap_path = os.path.join(outdir, f"{stem}.annotated.docx")
            n = annotate_docx(src, ap_path, findings)
        elif fmt == "hwpx":
            ap_path = os.path.join(outdir, f"{stem}.annotated.hwpx")
            n = annotate_hwpx(src, ap_path, findings)
        else:  # hwp / hwp-preview
            ap_path = os.path.join(outdir, f"{stem}.annotated.txt")
            n = annotate_hwp_fallback(segs, ap_path, findings)
        produced.append(ap_path)
        findings.notes.append(f"annotated {n} location(s) -> {ap_path}")

    print(json.dumps({
        "format": fmt, "style": style,
        "citations": findings.n_citations, "entries": findings.n_entries,
        "phantom": len(findings.phantom), "orphan": len(findings.orphan),
        "duplicate": len(findings.duplicate),
        "multiloc": len(findings.multiloc),
        "numbering_issues": len(findings.numbering),
        "outputs": produced,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
